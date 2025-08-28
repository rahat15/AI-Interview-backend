import os
from typing import Optional
import pypdf
from docx import Document


async def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return await extract_text_from_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return await extract_text_from_docx(file_path)
    elif file_extension in ['.txt', '.md']:
        return await extract_text_from_text(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


async def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


async def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word document"""
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Failed to extract text from Word document: {str(e)}")


async def extract_text_from_text(file_path: str) -> str:
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
        except Exception as e:
            raise Exception(f"Failed to read text file with any encoding: {str(e)}")
    
    except Exception as e:
        raise Exception(f"Failed to read text file: {str(e)}")


def get_file_metadata(file_path: str) -> dict:
    """Get basic file metadata"""
    if not os.path.exists(file_path):
        return {}
    
    stat = os.stat(file_path)
    return {
        "file_size": stat.st_size,
        "created_time": stat.st_ctime,
        "modified_time": stat.st_mtime,
        "file_extension": os.path.splitext(file_path)[1].lower(),
        "filename": os.path.basename(file_path)
    }
