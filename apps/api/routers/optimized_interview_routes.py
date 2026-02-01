"""
Optimized Interview API Routes with Streaming Support
Uses the new optimized engine for better performance
"""

from fastapi import APIRouter, HTTPException, File, UploadFile, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, Dict, Any, AsyncIterator
import json
import logging
import time

from interview.optimized_engine import optimized_engine
from interview.performance import monitor, log_performance_tips

logger = logging.getLogger(__name__)

router = APIRouter()


# ============================================================================
# Request/Response Models
# ============================================================================

class OptimizedStartRequest(BaseModel):
    user_id: str
    session_id: str
    role: str
    company: str
    cv_text: str
    jd_text: str


class OptimizedAnswerRequest(BaseModel):
    session_id: str
    answer: str
    voice_metrics: Optional[Dict[str, Any]] = None


# ============================================================================
# Helper Functions
# ============================================================================

async def fetch_content_from_db(content_id: str, collection: str) -> Optional[str]:
    """Fetch CV or JD content from MongoDB"""
    try:
        from bson import ObjectId
        from core.db import get_database
        
        db = await get_database()
        
        # Try different ID formats
        doc = None
        if ObjectId.is_valid(content_id):
            doc = await db[collection].find_one({"_id": ObjectId(content_id)})
        
        if not doc:
            doc = await db[collection].find_one({"_id": content_id})
        
        if doc:
            # Extract text content
            return doc.get("text", doc.get("content", ""))
        
        return None
    except Exception as e:
        logger.error(f"Error fetching {collection} {content_id}: {e}")
        return None


# ============================================================================
# Optimized Routes
# ============================================================================

@router.post("/v2/start")
async def start_optimized_session(req: OptimizedStartRequest):
    """
    Start optimized interview session with:
    - Fast context analysis with caching
    - Async operations
    - MongoDB persistence
    """
    try:
        result = await optimized_engine.create_session(
            session_id=req.session_id,
            user_id=req.user_id,
            role=req.role,
            company=req.company,
            cv_text=req.cv_text,
            jd_text=req.jd_text
        )
        
        return result
    except Exception as e:
        logger.error(f"Error starting session: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/v2/start-with-ids")
async def start_session_with_ids(
    user_id: str = Form(...),
    session_id: str = Form(...),
    role: str = Form(...),
    company: str = Form(...),
    cv_id: Optional[str] = Form(None),
    jd_id: Optional[str] = Form(None),
    cv_file: Optional[UploadFile] = File(None),
    jd_file: Optional[UploadFile] = File(None),
    cv_text: Optional[str] = Form(None),
    jd_text: Optional[str] = Form(None)
):
    """
    Start session with CV/JD files, IDs, or direct text
    Priority: files > IDs > text
    Automatically extracts text from files and fetches from MongoDB if IDs provided
    """
    try:
        # Fetch CV content (priority: file > id > text)
        final_cv = cv_text or ""
        
        if cv_file:
            # Extract text from uploaded file
            cv_content = await cv_file.read()
            file_ext = cv_file.filename.split('.')[-1].lower() if cv_file.filename else 'txt'
            
            if file_ext == 'txt':
                final_cv = cv_content.decode('utf-8')
            elif file_ext == 'pdf':
                # Extract text from PDF
                try:
                    from pypdf import PdfReader
                    import io
                    pdf_reader = PdfReader(io.BytesIO(cv_content))
                    final_cv = "\n".join([page.extract_text() for page in pdf_reader.pages])
                except Exception as e:
                    logger.error(f"PDF extraction error: {e}")
                    raise HTTPException(status_code=400, detail="Failed to extract text from PDF")
            elif file_ext in ['docx', 'doc']:
                # Extract text from DOCX
                try:
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(cv_content))
                    final_cv = "\n".join([para.text for para in doc.paragraphs])
                except Exception as e:
                    logger.error(f"DOCX extraction error: {e}")
                    raise HTTPException(status_code=400, detail="Failed to extract text from DOCX")
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported CV file format: {file_ext}")
            
            logger.info(f"✅ Extracted CV from file: {len(final_cv)} characters")
        elif cv_id:
            fetched_cv = await fetch_content_from_db(cv_id, "resumes")
            if fetched_cv:
                final_cv = fetched_cv
            else:
                raise HTTPException(status_code=404, detail=f"Resume {cv_id} not found")
        
        # Fetch JD content (priority: file > id > text)
        final_jd = jd_text or ""
        
        if jd_file:
            # Extract text from uploaded file
            jd_content = await jd_file.read()
            file_ext = jd_file.filename.split('.')[-1].lower() if jd_file.filename else 'txt'
            
            if file_ext == 'txt':
                final_jd = jd_content.decode('utf-8')
            elif file_ext == 'pdf':
                try:
                    from pypdf import PdfReader
                    import io
                    pdf_reader = PdfReader(io.BytesIO(jd_content))
                    final_jd = "\n".join([page.extract_text() for page in pdf_reader.pages])
                except Exception as e:
                    logger.error(f"PDF extraction error: {e}")
                    raise HTTPException(status_code=400, detail="Failed to extract text from JD PDF")
            elif file_ext in ['docx', 'doc']:
                try:
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(jd_content))
                    final_jd = "\n".join([para.text for para in doc.paragraphs])
                except Exception as e:
                    logger.error(f"DOCX extraction error: {e}")
                    raise HTTPException(status_code=400, detail="Failed to extract text from JD DOCX")
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported JD file format: {file_ext}")
            
            logger.info(f"✅ Extracted JD from file: {len(final_jd)} characters")
        elif jd_id:
            fetched_jd = await fetch_content_from_db(jd_id, "jobdescriptions")
            if fetched_jd:
                final_jd = fetched_jd
            else:
                raise HTTPException(status_code=404, detail=f"JD {jd_id} not found")
        
        if not final_cv or not final_jd:
            raise HTTPException(status_code=400, detail="Both CV and JD content required")
        
        result = await optimized_engine.create_session(
            session_id=session_id,
            user_id=user_id,
            role=role,
            company=company,
            cv_text=final_cv,
            jd_text=final_jd
        )
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error starting session: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/v2/answer")
async def submit_optimized_answer(
    session_id: str = Form(...),
    audio_file: Optional[UploadFile] = File(None),
    video_file: Optional[UploadFile] = File(None)
):
    """
    Submit answer with audio/video support
    - Concurrent transcription and analysis
    - Fast evaluation with caching
    """
    try:
        answer_text = ""
        voice_metrics = None
        
        # Process audio if provided
        if audio_file:
            audio_data = await audio_file.read()
            
            if len(audio_data) > 100:
                # Transcribe audio
                from interview.speech_to_text import speech_converter
                answer_text = speech_converter.convert_audio_to_text(audio_data)
                
                # Analyze voice (async)
                try:
                    from interview.voice_analyzer import voice_analyzer
                    voice_metrics = voice_analyzer.analyze_voice(audio_data, answer_text)
                except Exception as e:
                    logger.warning(f"Voice analysis failed: {e}")
        
        # Extract audio from video if no audio file
        if not answer_text and video_file:
            video_data = await video_file.read()
            
            # Extract and transcribe audio from video
            import tempfile
            import subprocess
            import os
            
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as video_tmp:
                    video_tmp.write(video_data)
                    video_path = video_tmp.name
                
                audio_path = video_path.replace('.mp4', '.wav')
                subprocess.run(
                    ['ffmpeg', '-i', video_path, '-vn', '-acodec', 'pcm_s16le', '-ar', '16000', '-ac', '1', audio_path, '-y'],
                    capture_output=True,
                    timeout=10
                )
                
                if os.path.exists(audio_path):
                    with open(audio_path, 'rb') as f:
                        audio_data = f.read()
                    
                    from interview.speech_to_text import speech_converter
                    answer_text = speech_converter.convert_audio_to_text(audio_data)
                    
                    os.remove(audio_path)
                
                os.remove(video_path)
            except Exception as e:
                logger.error(f"Video audio extraction failed: {e}")
        
        if not answer_text:
            answer_text = "No response detected"
        
        # Submit answer using optimized engine
        result = await optimized_engine.submit_answer(
            session_id=session_id,
            answer=answer_text,
            voice_metrics=voice_metrics
        )
        
        return result
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Error submitting answer: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/v2/stream/{session_id}")
async def stream_next_question(session_id: str):
    """
    Stream next question generation for real-time feel
    Returns Server-Sent Events (SSE)
    """
    async def event_generator() -> AsyncIterator[str]:
        try:
            async for chunk in optimized_engine.stream_question(session_id):
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            
            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            logger.error(f"Streaming error: {e}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
        }
    )


@router.get("/v2/state/{session_id}")
async def get_optimized_state(session_id: str):
    """Get current session state"""
    try:
        state = await optimized_engine.get_session_state(session_id)
        
        if not state:
            raise HTTPException(status_code=404, detail="Session not found")
        
        # Return serializable state
        return {
            "session_id": state.get("session_id"),
            "user_id": state.get("user_id"),
            "role": state.get("role"),
            "company": state.get("company"),
            "question_count": state.get("question_count", 0),
            "stage": state.get("stage", "intro"),
            "completed": state.get("completed", False),
            "messages": state.get("messages", []),
            "avg_response_time": sum(state.get("response_times", [])) / len(state.get("response_times", [])) if state.get("response_times") else 0
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting state: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/v2/performance/{session_id}")
async def get_performance_metrics(session_id: str):
    """Get performance metrics for session"""
    try:
        state = await optimized_engine.get_session_state(session_id)
        
        if not state:
            raise HTTPException(status_code=404, detail="Session not found")
        
        response_times = state.get("response_times", [])
        
        return {
            "session_id": session_id,
            "total_questions": state.get("question_count", 0),
            "response_times": {
                "min": min(response_times) if response_times else 0,
                "max": max(response_times) if response_times else 0,
                "avg": sum(response_times) / len(response_times) if response_times else 0,
                "all": response_times
            },
            "cache_status": "active" if state.get("cv_summary") else "not_cached"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting metrics: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/v2/complete/{session_id}")
async def complete_optimized_interview(session_id: str):
    """
    Complete interview and generate final evaluation
    Uses cached context for fast evaluation
    """
    try:
        state = await optimized_engine.get_session_state(session_id)
        
        if not state:
            raise HTTPException(status_code=404, detail="Session not found")
        
        # Generate comprehensive evaluation
        from interview.optimized_engine import AsyncAnswerEvaluator
        evaluator = AsyncAnswerEvaluator()
        
        messages = state.get("messages", [])
        
        # Collect all Q&A pairs
        qa_pairs = []
        for i in range(0, len(messages) - 1, 2):
            if i + 1 < len(messages):
                qa_pairs.append({
                    "question": messages[i].get("content"),
                    "answer": messages[i + 1].get("content"),
                    "evaluation": messages[i + 1].get("metadata", {}).get("evaluation")
                })
        
        # Calculate overall scores
        all_evaluations = [qa.get("evaluation") for qa in qa_pairs if qa.get("evaluation")]
        
        avg_clarity = sum(e.get("clarity", 0) for e in all_evaluations) / len(all_evaluations) if all_evaluations else 7
        avg_relevance = sum(e.get("relevance", 0) for e in all_evaluations) / len(all_evaluations) if all_evaluations else 7
        avg_depth = sum(e.get("depth", 0) for e in all_evaluations) / len(all_evaluations) if all_evaluations else 7
        
        overall_score = (avg_clarity + avg_relevance + avg_depth) / 3
        
        recommendation = "hire" if overall_score >= 8 else "maybe" if overall_score >= 6 else "no_hire"
        
        return {
            "session_id": session_id,
            "status": "completed",
            "total_questions": state.get("question_count", 0),
            "evaluation": {
                "overall_score": round(overall_score, 2),
                "recommendation": recommendation,
                "clarity": round(avg_clarity, 2),
                "relevance": round(avg_relevance, 2),
                "depth": round(avg_depth, 2)
            },
            "conversation": qa_pairs,
            "performance_metrics": {
                "avg_response_time": sum(state.get("response_times", [])) / len(state.get("response_times", [])) if state.get("response_times") else 0,
                "total_response_times": state.get("response_times", [])
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error completing interview: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/v2/metrics/global")
async def get_global_metrics():
    """Get global performance metrics across all sessions"""
    try:
        stats = monitor.get_stats()
        log_performance_tips()
        
        return {
            "status": "success",
            "metrics": stats,
            "timestamp": time.time()
        }
    except Exception as e:
        logger.error(f"Error getting global metrics: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/v2/metrics/reset")
async def reset_metrics():
    """Reset performance metrics (admin only)"""
    try:
        monitor.reset_metrics()
        return {"status": "success", "message": "Metrics reset successfully"}
    except Exception as e:
        logger.error(f"Error resetting metrics: {e}")
        raise HTTPException(status_code=500, detail=str(e))
